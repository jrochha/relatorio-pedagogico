from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file
from datetime import datetime
import json
import os
from pathlib import Path
from io import BytesIO

from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader


app = Flask(__name__, template_folder="templates", static_folder="static")

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
STATIC_DIR = BASE_DIR / "static"

LOGO_ESCOLA_PATH = STATIC_DIR / "logo_escola.png"
LOGO_ESTADO_PATH = STATIC_DIR / "logo_estado.png"

database_url = os.getenv("DATABASE_URL", f"sqlite:///{BASE_DIR / 'relatorios.db'}")
if database_url.startswith("postgres://"):
    database_url = database_url.replace("postgres://", "postgresql://", 1)

app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)


with open(DATA_DIR / "alunos.json", "r", encoding="utf-8") as f:
    estudantes = json.load(f)

with open(DATA_DIR / "config.json", "r", encoding="utf-8") as f:
    config = json.load(f)


class Registro(db.Model):
    __tablename__ = "registros"

    id = db.Column(db.Integer, primary_key=True)
    estudante_id = db.Column(db.Integer, nullable=False, index=True)

    data = db.Column(db.String(20), nullable=False)
    hora = db.Column(db.String(10), nullable=False)
    tipo_relatorio = db.Column(db.String(150), nullable=False)
    disciplina = db.Column(db.String(150), nullable=False)
    profissional = db.Column(db.String(150), nullable=False)

    dificuldades = db.Column(db.Text, default="[]")
    questoes_comportamentais = db.Column(db.Text, default="[]")
    intervencoes = db.Column(db.Text, default="[]")
    resposta_aluno = db.Column(db.String(250), default="")
    encaminhamentos = db.Column(db.Text, default="[]")
    observacao_livre = db.Column(db.Text, default="")

    assinatura_estudante = db.Column(db.String(150), default="")
    assinatura_profissional = db.Column(db.String(150), default="")
    assinatura_responsavel = db.Column(db.String(150), default="")

    relatorio_texto = db.Column(db.Text, nullable=False)
    criado_em = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)

    def lista(self, campo):
        valor = getattr(self, campo, "[]")
        try:
            return json.loads(valor)
        except Exception:
            return []


with app.app_context():
    db.create_all()


def buscar_estudante(estudante_id):
    return next((e for e in estudantes if e["id"] == estudante_id), None)


def buscar_registro(registro_id):
    return db.session.get(Registro, registro_id)


def valor_registro(registro, campo, padrao=""):
    if isinstance(registro, dict):
        return registro.get(campo, padrao)
    return getattr(registro, campo, padrao)


def lista_registro(registro, campo):
    if isinstance(registro, dict):
        return registro.get(campo, [])

    valor = getattr(registro, campo, "[]")
    try:
        return json.loads(valor)
    except Exception:
        return []


def formatar_lista(lista):
    if not lista:
        return ""
    if len(lista) == 1:
        return lista[0].lower()
    if len(lista) == 2:
        return f"{lista[0].lower()} e {lista[1].lower()}"
    return ", ".join(item.lower() for item in lista[:-1]) + f" e {lista[-1].lower()}"


def formatar_data_br(data_iso):
    if not data_iso:
        return datetime.now().strftime("%d/%m/%Y")
    try:
        return datetime.strptime(data_iso, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return data_iso


def gerar_relatorio(estudante, registro):
    tipo_relatorio = valor_registro(registro, "tipo_relatorio")
    disciplina = valor_registro(registro, "disciplina")

    dificuldades = formatar_lista(lista_registro(registro, "dificuldades"))
    comportamentos = formatar_lista(lista_registro(registro, "questoes_comportamentais"))
    intervencoes = formatar_lista(lista_registro(registro, "intervencoes"))
    encaminhamentos = formatar_lista(lista_registro(registro, "encaminhamentos"))

    resposta = valor_registro(registro, "resposta_aluno", "não informada").lower()
    observacao = valor_registro(registro, "observacao_livre", "").strip()

    partes = [
        f"{tipo_relatorio} referente ao(à) estudante {estudante['nome']}, da turma {estudante['turma']}, "
        f"registrado em {valor_registro(registro, 'data')} às {valor_registro(registro, 'hora')}."
    ]

    if tipo_relatorio == "Relatório de Ocorrência Disciplinar":
        partes.append(f"O registro foi realizado no contexto da disciplina de {disciplina}.")
        if lista_registro(registro, "questoes_comportamentais"):
            partes.append(f"Foram observadas as seguintes ocorrências comportamentais: {comportamentos}.")
        if lista_registro(registro, "intervencoes"):
            partes.append(f"Diante da situação, foram adotadas as seguintes intervenções: {intervencoes}.")
        partes.append(f"Após a intervenção, verificou-se que o(a) estudante {resposta}.")
        if lista_registro(registro, "encaminhamentos"):
            partes.append(f"Como orientação, recomenda-se {encaminhamentos}.")
        if observacao:
            partes.append(f"Observação complementar: {observacao}")
    else:
        partes.append(f"O registro foi realizado no contexto da disciplina de {disciplina}.")
        if lista_registro(registro, "dificuldades"):
            partes.append(f"No âmbito da aprendizagem, foram observadas dificuldades relacionadas a {dificuldades}.")
        if lista_registro(registro, "intervencoes"):
            partes.append(f"Como intervenções pedagógicas, foram realizadas ações como {intervencoes}.")
        partes.append(f"Após as intervenções, verificou-se que o(a) estudante {resposta}.")
        if lista_registro(registro, "encaminhamentos"):
            partes.append(f"Como orientação, recomenda-se {encaminhamentos}.")
        if observacao:
            partes.append(f"Observação complementar: {observacao}")

    return " ".join(partes)


def nome_arquivo_base(estudante, registro):
    base = f"{estudante['nome']}_{valor_registro(registro, 'tipo_relatorio')}_{valor_registro(registro, 'data')}"
    return base.replace("/", "-").replace(" ", "_")


def texto_completo_exportacao(estudante, registro):
    linhas = [
        "ESTADO DO PARANÁ",
        "SECRETARIA DE ESTADO DA EDUCAÇÃO",
        "PARANÁ INTEGRAL",
        "ESCOLA ESTADUAL PADRE MANUEL DA NÓBREGA",
        "",
        "Nº ________/2026",
        valor_registro(registro, "tipo_relatorio").upper(),
        "",
        f"Estudante: {estudante['nome']}",
        f"Turma: {estudante['turma']}",
        f"Responsável: {estudante['responsavel']}",
        f"Data: {valor_registro(registro, 'data')}",
        f"Hora: {valor_registro(registro, 'hora')}",
        f"Profissional: {valor_registro(registro, 'profissional')}",
        f"Disciplina: {valor_registro(registro, 'disciplina')}",
        "",
        valor_registro(registro, "relatorio_texto"),
        "",
        "",
        "________________________________________",
        f"Assinatura do(a) estudante: {valor_registro(registro, 'assinatura_estudante')}",
        "",
        "________________________________________",
        f"Assinatura do(a) professor(a)/pedagogo(a): {valor_registro(registro, 'assinatura_profissional')}",
        "",
        "________________________________________",
        f"Assinatura do responsável: {valor_registro(registro, 'assinatura_responsavel')}",
    ]
    return "\n".join(linhas)


def inserir_logos_doc(document):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_ESTADO_PATH.exists():
        p_left.add_run().add_picture(str(LOGO_ESTADO_PATH), width=Inches(1.6))

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_ESCOLA_PATH.exists():
        p_right.add_run().add_picture(str(LOGO_ESCOLA_PATH), width=Inches(1.0))


def gerar_docx(estudante, registro):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    inserir_logos_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "ESTADO DO PARANÁ\n"
        "SECRETARIA DE ESTADO DA EDUCAÇÃO\n"
        "PARANÁ INTEGRAL\n"
        "ESCOLA ESTADUAL PADRE MANUEL DA NÓBREGA"
    )
    r.bold = True
    r.font.size = Pt(12)

    num = doc.add_paragraph()
    num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_run = num.add_run("Nº ________/2026")
    num_run.bold = True
    num_run.font.size = Pt(11)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.add_run(valor_registro(registro, "tipo_relatorio").upper())
    titulo_run.bold = True
    titulo_run.font.size = Pt(12)

    doc.add_paragraph(f"Estudante: {estudante['nome']}")
    doc.add_paragraph(f"Turma: {estudante['turma']}")
    doc.add_paragraph(f"Responsável: {estudante['responsavel']}")
    doc.add_paragraph(f"Data: {valor_registro(registro, 'data')}    Hora: {valor_registro(registro, 'hora')}")
    doc.add_paragraph(f"Profissional: {valor_registro(registro, 'profissional')}")
    doc.add_paragraph(f"Disciplina: {valor_registro(registro, 'disciplina')}")
    doc.add_paragraph("")

    corpo = doc.add_paragraph(valor_registro(registro, "relatorio_texto"))
    corpo.paragraph_format.line_spacing = 1.5

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("________________________________________")
    doc.add_paragraph(f"Assinatura do(a) estudante: {valor_registro(registro, 'assinatura_estudante')}")
    doc.add_paragraph("")
    doc.add_paragraph("________________________________________")
    doc.add_paragraph(f"Assinatura do(a) professor(a)/pedagogo(a): {valor_registro(registro, 'assinatura_profissional')}")
    doc.add_paragraph("")
    doc.add_paragraph("________________________________________")
    doc.add_paragraph(f"Assinatura do responsável: {valor_registro(registro, 'assinatura_responsavel')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def draw_wrapped_text(pdf, text, x, y, max_width, font_name="Helvetica", font_size=11, leading=16):
    words = text.split()
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if stringWidth(test, font_name, font_size) <= max_width:
            line = test
        else:
            pdf.drawString(x, y, line)
            y -= leading
            line = word
    if line:
        pdf.drawString(x, y, line)
        y -= leading
    return y


def draw_header_pdf(pdf, width, height):
    y = height - 55

    if LOGO_ESTADO_PATH.exists():
        try:
            pdf.drawImage(
                ImageReader(str(LOGO_ESTADO_PATH)),
                40,
                y - 55,
                width=105,
                height=55,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

    if LOGO_ESCOLA_PATH.exists():
        try:
            pdf.drawImage(
                ImageReader(str(LOGO_ESCOLA_PATH)),
                width - 85,
                y - 52,
                width=42,
                height=42,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawCentredString(width / 2, y, "ESTADO DO PARANÁ")
    pdf.drawCentredString(width / 2, y - 14, "SECRETARIA DE ESTADO DA EDUCAÇÃO")
    pdf.drawCentredString(width / 2, y - 28, "PARANÁ INTEGRAL")
    pdf.drawCentredString(width / 2, y - 42, "ESCOLA ESTADUAL PADRE MANUEL DA NÓBREGA")
    pdf.line(40, y - 58, width - 40, y - 58)

    return y - 78


def gerar_pdf(estudante, registro):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = draw_header_pdf(pdf, width, height)

    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawRightString(width - 45, y, "Nº ________/2026")
    y -= 24

    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawCentredString(width / 2, y, valor_registro(registro, "tipo_relatorio").upper())
    y -= 28

    pdf.setFont("Helvetica", 11)
    for campo in [
        f"Estudante: {estudante['nome']}",
        f"Turma: {estudante['turma']}",
        f"Responsável: {estudante['responsavel']}",
        f"Data: {valor_registro(registro, 'data')}    Hora: {valor_registro(registro, 'hora')}",
        f"Profissional: {valor_registro(registro, 'profissional')}",
        f"Disciplina: {valor_registro(registro, 'disciplina')}",
    ]:
        pdf.drawString(45, y, campo)
        y -= 18

    y -= 8
    y = draw_wrapped_text(pdf, valor_registro(registro, "relatorio_texto"), 45, y, width - 90, leading=17)

    y -= 20
    for titulo, assinatura in [
        ("Assinatura do(a) estudante:", valor_registro(registro, "assinatura_estudante")),
        ("Assinatura do(a) professor(a)/pedagogo(a):", valor_registro(registro, "assinatura_profissional")),
        ("Assinatura do responsável:", valor_registro(registro, "assinatura_responsavel")),
    ]:
        if y < 120:
            pdf.showPage()
            y = draw_header_pdf(pdf, width, height) - 20
            pdf.setFont("Helvetica", 11)

        pdf.line(45, y, 260, y)
        y -= 16
        pdf.drawString(45, y, f"{titulo} {assinatura}")
        y -= 34

    pdf.save()
    buffer.seek(0)
    return buffer


@app.route("/")
def index():
    turma = request.args.get("turma", "").strip()
    busca = request.args.get("q", "").strip().lower()

    estudantes_filtrados = estudantes

    if turma:
        estudantes_filtrados = [e for e in estudantes_filtrados if e["turma"] == turma]

    if busca:
        estudantes_filtrados = [
            e for e in estudantes_filtrados
            if busca in e["nome"].lower() or busca in e["responsavel"].lower()
        ]

    total_registros = db.session.query(func.count(Registro.id)).scalar() or 0

    return render_template(
        "index.html",
        estudantes=estudantes_filtrados,
        turmas=sorted({e["turma"] for e in estudantes}),
        turma_atual=turma,
        busca=busca,
        total_estudantes=len(estudantes),
        total_registros=total_registros,
    )


@app.route("/api/estudante/<int:estudante_id>")
def api_estudante(estudante_id):
    estudante = buscar_estudante(estudante_id)
    if not estudante:
        return jsonify({"erro": "Estudante não encontrado"}), 404

    total = Registro.query.filter_by(estudante_id=estudante_id).count()
    estudante_com_historico = dict(estudante)
    estudante_com_historico["total_registros"] = total
    return jsonify(estudante_com_historico)


@app.route("/novo_registro", methods=["GET", "POST"])
def novo_registro():
    if request.method == "POST":
        estudante_id = int(request.form["estudante_id"])
        estudante = buscar_estudante(estudante_id)

        if not estudante:
            return "Estudante não encontrado.", 404

        registro_dict = {
            "estudante_id": estudante_id,
            "data": formatar_data_br(request.form.get("data_registro", "").strip()),
            "hora": request.form.get("hora_registro", "").strip() or datetime.now().strftime("%H:%M"),
            "tipo_relatorio": request.form["tipo_relatorio"],
            "disciplina": request.form["disciplina"],
            "profissional": request.form["profissional"].strip(),
            "dificuldades": request.form.getlist("dificuldades"),
            "questoes_comportamentais": request.form.getlist("questoes_comportamentais"),
            "intervencoes": request.form.getlist("intervencoes"),
            "resposta_aluno": request.form.get("resposta_aluno", ""),
            "encaminhamentos": request.form.getlist("encaminhamentos"),
            "observacao_livre": request.form.get("observacao_livre", "").strip(),
            "assinatura_estudante": request.form.get("assinatura_estudante", "").strip(),
            "assinatura_profissional": request.form.get("assinatura_profissional", "").strip(),
            "assinatura_responsavel": request.form.get("assinatura_responsavel", "").strip(),
        }

        registro_dict["relatorio_texto"] = gerar_relatorio(estudante, registro_dict)

        novo = Registro(
            estudante_id=registro_dict["estudante_id"],
            data=registro_dict["data"],
            hora=registro_dict["hora"],
            tipo_relatorio=registro_dict["tipo_relatorio"],
            disciplina=registro_dict["disciplina"],
            profissional=registro_dict["profissional"],
            dificuldades=json.dumps(registro_dict["dificuldades"], ensure_ascii=False),
            questoes_comportamentais=json.dumps(registro_dict["questoes_comportamentais"], ensure_ascii=False),
            intervencoes=json.dumps(registro_dict["intervencoes"], ensure_ascii=False),
            resposta_aluno=registro_dict["resposta_aluno"],
            encaminhamentos=json.dumps(registro_dict["encaminhamentos"], ensure_ascii=False),
            observacao_livre=registro_dict["observacao_livre"],
            assinatura_estudante=registro_dict["assinatura_estudante"],
            assinatura_profissional=registro_dict["assinatura_profissional"],
            assinatura_responsavel=registro_dict["assinatura_responsavel"],
            relatorio_texto=registro_dict["relatorio_texto"],
        )

        db.session.add(novo)
        db.session.commit()

        return redirect(url_for("relatorio_estudante", estudante_id=estudante_id))

    return render_template(
        "novo_registro.html",
        estudantes=estudantes,
        turmas=sorted({e["turma"] for e in estudantes}),
        config=config,
        data_atual=datetime.now().strftime("%Y-%m-%d"),
        hora_atual=datetime.now().strftime("%H:%M"),
    )


@app.route("/relatorio/<int:estudante_id>")
def relatorio_estudante(estudante_id):
    estudante = buscar_estudante(estudante_id)

    if not estudante:
        return "Estudante não encontrado.", 404

    registros_estudante = (
        Registro.query
        .filter_by(estudante_id=estudante_id)
        .order_by(Registro.criado_em.desc(), Registro.id.desc())
        .all()
    )

    return render_template(
        "relatorio.html",
        estudante=estudante,
        registros=registros_estudante,
    )


@app.route("/historico/<int:estudante_id>")
def historico_estudante(estudante_id):
    return relatorio_estudante(estudante_id)


@app.route("/download/<int:registro_id>/<string:formato>")
def download_relatorio(registro_id, formato):
    registro = buscar_registro(registro_id)

    if not registro:
        return "Registro não encontrado.", 404

    estudante = buscar_estudante(registro.estudante_id)

    if not estudante:
        return "Estudante não encontrado.", 404

    nome_base = nome_arquivo_base(estudante, registro)

    if formato == "txt":
        conteudo = texto_completo_exportacao(estudante, registro)
        buffer = BytesIO(conteudo.encode("utf-8"))
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{nome_base}.txt",
            mimetype="text/plain; charset=utf-8",
        )

    if formato == "docx":
        return send_file(
            gerar_docx(estudante, registro),
            as_attachment=True,
            download_name=f"{nome_base}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if formato == "pdf":
        return send_file(
            gerar_pdf(estudante, registro),
            as_attachment=True,
            download_name=f"{nome_base}.pdf",
            mimetype="application/pdf",
        )

    return "Formato inválido.", 400


if __name__ == "__main__":
    app.run(debug=True)
